from pathlib import Path

from app.ingestion.extractors import DocumentExtractor


def test_extract_html_sections(tmp_path):
    html_content = """
    <html>
      <head><title>Remote Work Policy</title></head>
      <body>
        <h1>Remote Work Policy</h1>
        <p>Effective Date: January 1, 2024</p>
        <h2>Eligibility</h2>
        <p>All full-time employees may request remote work.</p>
      </body>
    </html>
    """
    html_path = tmp_path / "remote_policy.html"
    html_path.write_text(html_content.strip(), encoding="utf-8")

    extractor = DocumentExtractor()
    result = extractor.extract(html_path)

    assert result.documents
    first_doc = result.documents[0]
    assert first_doc.metadata["doc_type"] == "html"
    assert first_doc.metadata["policy_name"].startswith("Remote Work")
    assert first_doc.metadata.get("effective_date") == "2024-01-01"
    viewer_path = Path(first_doc.metadata["viewer_path"])
    assert viewer_path.exists()
    assert first_doc.metadata["viewer_page"] == 1


def test_extract_docx_sections(tmp_path):
    docx_path = tmp_path / "travel_policy.docx"

    from docx import Document  # type: ignore

    doc = Document()
    doc.add_heading("Travel Policy", level=1)
    doc.add_paragraph("Effective Date: March 3, 2023")
    doc.add_paragraph("All travel must be pre-approved by the finance team.")
    doc.save(str(docx_path))

    extractor = DocumentExtractor()
    result = extractor.extract(docx_path)

    assert result.documents
    doc_meta = result.documents[0].metadata
    assert doc_meta["doc_type"] == "docx"
    assert doc_meta["policy_name"].startswith("Travel Policy")
    assert doc_meta.get("effective_date") == "2023-03-03"
    viewer_path = Path(doc_meta["viewer_path"])
    assert viewer_path.exists()
    assert doc_meta["viewer_page"] == 1
