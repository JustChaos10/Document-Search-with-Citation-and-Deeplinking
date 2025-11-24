from __future__ import annotations

import logging
import re
import textwrap
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, Iterable, List, Optional, Sequence, Tuple

from langchain_core.documents import Document
from pypdf import PdfReader

from app.nlp import detect_language

if TYPE_CHECKING:
    from app.audio.transcription import SpeechToTextService, TranscriptionResult

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    documents: List[Document]


class DocumentExtractor:
    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".mp3",
        ".wav",
        ".m4a",
        ".mp4",
        ".flac",
        ".html",
        ".htm",
        ".docx",
    }
    AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a", ".mp4", ".flac"}
    HTML_EXTENSIONS = {".html", ".htm"}

    def __init__(self, stt_service: Optional["SpeechToTextService"] = None) -> None:
        self._stt_service = stt_service
        try:  # Local import so optional dependency remains optional
            from bs4 import BeautifulSoup  # noqa: F401

            self._has_bs4 = True
        except ImportError:  # pragma: no cover - optional dependency
            self._has_bs4 = False
        try:
            from docx import Document as DocxDocument  # noqa: F401

            self._has_docx = True
        except ImportError:  # pragma: no cover - optional dependency
            self._has_docx = False
        try:
            from reportlab.pdfgen import canvas  # noqa: F401

            self._has_reportlab = True
        except ImportError:  # pragma: no cover - optional dependency
            self._has_reportlab = False


    def extract(self, file_path: Path) -> ExtractionResult:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {suffix}")

        logger.info("extract.start", extra={"path": str(file_path), "extension": suffix})

        loader_map = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
        }
        if suffix in self.AUDIO_EXTENSIONS:
            loader_map[suffix] = self._load_audio
        if suffix in self.HTML_EXTENSIONS:
            loader_map[suffix] = self._load_html

        loader = loader_map[suffix]
        documents = list(loader(file_path))

        logger.info(
            "extract.complete",
            extra={"path": str(file_path), "extension": suffix, "sections": len(documents)},
        )
        return ExtractionResult(documents=documents)

    def _base_metadata(self, file_path: Path) -> dict:
        return {
            "source": file_path.name,
            "path": str(file_path.resolve()),
        }

    def _normalize_text(self, text: str) -> str:
        if not text:
            return ""

        # Replace non-breaking spaces but preserve Arabic text structure
        normalized = text.replace("\u00a0", " ").strip()
        if not normalized:
            return ""

        # Detect if text contains Arabic
        has_arabic = bool(re.search(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]", normalized))

        # Remove problematic characters that might come from PDF extraction
        # Remove zero-width characters and other invisible Unicode that can corrupt Arabic
        normalized = re.sub(r"[\u200B-\u200D\uFEFF]", "", normalized)  # Zero-width chars

        # Normalize whitespace more carefully for Arabic
        if has_arabic:
            # Preserve paragraph structure better for Arabic
            paragraphs = re.split(r"\n{2,}", normalized)
            cleaned_paragraphs = []
            for paragraph in paragraphs:
                # Normalize spaces but preserve Arabic word boundaries
                # Only collapse multiple spaces, preserve single spaces
                flattened = re.sub(r"[ \t]+", " ", paragraph.strip())
                # Remove standalone ASCII dashes that might be artifacts
                flattened = re.sub(r"\s+-\s+", " ", flattened)
                if flattened:
                    cleaned_paragraphs.append(flattened)
            return "\n\n".join(cleaned_paragraphs)
        else:
            # For non-Arabic, use existing logic
            paragraphs = re.split(r"\n{2,}", normalized)
            cleaned_paragraphs = []
            for paragraph in paragraphs:
                flattened = re.sub(r"\s+", " ", paragraph.strip())
                if flattened:
                    cleaned_paragraphs.append(flattened)
            return "\n\n".join(cleaned_paragraphs)

    def _load_pdf(self, file_path: Path) -> Iterable[Document]:
        reader = PdfReader(str(file_path))
        for index, page in enumerate(reader.pages, start=1):
            # Try layout mode first for better Arabic/RTL support
            try:
                text = page.extract_text(extraction_mode="layout") or ""
            except (TypeError, AttributeError, ValueError):
                # Fallback for older pypdf versions or if layout mode fails
                try:
                    text = page.extract_text() or ""
                except Exception as e:
                    logger.warning(
                        "pdf.extract_failed",
                        extra={
                            "page": index,
                            "file": str(file_path),
                            "error": str(e)
                        }
                    )
                    continue

            # Validate Arabic text integrity (check for obvious corruption)
            if self._is_arabic_corrupted(text):
                logger.warning(
                    "pdf.arabic_corruption_detected",
                    extra={
                        "page": index,
                        "file": str(file_path)
                    }
                )

            text = self._normalize_text(text)
            if not text:
                continue

            # Detect language at page level before chunking
            detected_lang = detect_language(text)

            metadata = {
                **self._base_metadata(file_path),
                "page": index,
                "doc_type": "pdf",
                "original_doc_type": "pdf",
                "language": detected_lang if detected_lang else "en",  # Default to English
                "viewer_path": str(file_path.resolve()),
                "viewer_page": index,
            }
            metadata.update(self._infer_policy_metadata(text, fallback_title=metadata["source"]))
            yield Document(page_content=text, metadata=metadata)

    def _load_audio(self, file_path: Path) -> Iterable[Document]:
        if not self._stt_service:
            raise RuntimeError(
                "Audio transcription requested but SpeechToTextService is not configured."
            )
        transcription = self._stt_service.transcribe(file_path)
        if not transcription.text:
            logger.warning(
                "extract.audio.empty",
                extra={"path": str(file_path)},
            )
            return []

        language = transcription.language or detect_language(transcription.text) or "en"
        metadata = {
            **self._base_metadata(file_path),
            "doc_type": "audio",
            "language": language,
            "duration_seconds": transcription.duration,
            "segment_count": len(transcription.segments),
            "transcription_model": getattr(self._stt_service, "model_name", None),
        }
        yield Document(page_content=transcription.text, metadata=metadata)

    def _load_html(self, file_path: Path) -> Iterable[Document]:
        if not self._has_bs4:
            raise ImportError(
                "beautifulsoup4 is required to ingest HTML documents. "
                "Install it via `pip install beautifulsoup4`."
            )
        from bs4 import BeautifulSoup  # type: ignore

        raw = file_path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(raw, "html.parser")

        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        title = (soup.title.string.strip() if soup.title and soup.title.string else file_path.stem)
        raw_sections = self._split_html_sections(soup)
        if not raw_sections:
            full_text = soup.get_text(separator="\n")
            raw_sections = [(title, full_text)]

        sections: List[Tuple[Optional[str], str]] = []
        for heading, body_text in raw_sections:
            text = self._normalize_text(body_text)
            if text:
                sections.append((heading or title, text))

        if not sections:
            return

        surrogate_path = self._create_surrogate_pdf(file_path, sections)
        surrogate_resolved = surrogate_path.resolve()

        for idx, (heading, text) in enumerate(sections, start=1):
            inferred_lang = detect_language(text) or "en"
            metadata = {
                **self._base_metadata(file_path),
                "doc_type": "html",
                "original_doc_type": "html",
                "language": inferred_lang,
                "section_heading": heading or title,
                "section_index": idx,
                "policy_name": title,
                "viewer_path": str(surrogate_resolved),
                "viewer_page": idx,
                "page": idx,
            }
            metadata.update(self._infer_policy_metadata(text, fallback_title=title, heading=heading))
            yield Document(page_content=text, metadata=metadata)

    def _load_docx(self, file_path: Path) -> Iterable[Document]:
        if not self._has_docx:
            raise ImportError(
                "python-docx is required to ingest DOCX documents. Install it via `pip install python-docx`."
            )
        from docx import Document as DocxDocument  # type: ignore

        docx_doc = DocxDocument(str(file_path))
        title = docx_doc.core_properties.title or file_path.stem

        sections: List[Tuple[Optional[str], str]] = []
        current_heading = None
        current_buffer: List[str] = []

        def flush_section():
            if not current_buffer:
                return
            text = "\n".join(current_buffer).strip()
            if not text:
                return
            sections.append((current_heading, text))

        for paragraph in docx_doc.paragraphs:
            style_name = (paragraph.style.name if paragraph.style else "") or ""
            text = paragraph.text.strip()
            if not text:
                continue
            if style_name.lower().startswith("heading"):
                flush_section()
                current_heading = text
                current_buffer = []
            else:
                current_buffer.append(text)
        flush_section()

        if not sections:
            full_text = "\n".join(p.text for p in docx_doc.paragraphs if p.text.strip())
            sections = [(title, full_text)]

        prepared_sections: List[Tuple[Optional[str], str]] = []
        for heading, body in sections:
            text = self._normalize_text(body)
            if text:
                prepared_sections.append((heading or title, text))

        if not prepared_sections:
            return

        surrogate_path = self._create_surrogate_pdf(file_path, prepared_sections)
        surrogate_resolved = surrogate_path.resolve()

        for idx, (heading, text) in enumerate(prepared_sections, start=1):
            lang = detect_language(text) or "en"
            metadata = {
                **self._base_metadata(file_path),
                "doc_type": "docx",
                "original_doc_type": "docx",
                "language": lang,
                "section_heading": heading or title,
                "section_index": idx,
                "policy_name": title,
                "viewer_path": str(surrogate_resolved),
                "viewer_page": idx,
                "page": idx,
            }
            metadata.update(self._infer_policy_metadata(text, fallback_title=title, heading=heading))
            yield Document(page_content=text, metadata=metadata)

    def _split_html_sections(self, soup) -> List[Tuple[Optional[str], str]]:
        """Derive sections from HTML by walking heading tags."""
        sections: List[Tuple[Optional[str], str]] = []
        current_heading: Optional[str] = None
        current_chunks: List[str] = []
        for element in soup.body.descendants if soup.body else soup.descendants:
            if getattr(element, "name", None) and element.name.lower() in {"h1", "h2", "h3"}:
                heading_text = element.get_text(separator=" ", strip=True)
                if current_chunks:
                    sections.append((current_heading, "\n".join(current_chunks)))
                    current_chunks = []
                current_heading = heading_text
            elif isinstance(element, str):
                text = element.strip()
                if text:
                    current_chunks.append(text)
        if current_chunks:
            sections.append((current_heading, "\n".join(current_chunks)))
        return sections

    def _create_surrogate_pdf(
        self,
        file_path: Path,
        sections: Sequence[Tuple[Optional[str], str]],
    ) -> Path:
        if not self._has_reportlab:
            raise ImportError(
                "reportlab is required to render DOCX/HTML documents for the inline viewer. "
                "Install it via `pip install reportlab`."
            )
        from reportlab.lib.pagesizes import LETTER  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore

        converted_path = file_path.with_name(f"{file_path.stem}.converted.pdf")
        converted_path.parent.mkdir(parents=True, exist_ok=True)

        page_width, page_height = LETTER
        margin_x = 72
        margin_y = 72
        max_width = page_width - 2 * margin_x

        pdf = canvas.Canvas(str(converted_path), pagesize=LETTER)

        for heading, body in sections:
            text_obj = pdf.beginText(margin_x, page_height - margin_y)
            if heading:
                text_obj.setFont("Helvetica-Bold", 14)
                text_obj.textLine(heading.strip())
                text_obj.moveCursor(0, 14)
            text_obj.setFont("Helvetica", 11)
            for line in self._wrap_text(body, max_width):
                text_obj.textLine(line)
            pdf.drawText(text_obj)
            pdf.showPage()

        if not sections:
            pdf.drawString(margin_x, page_height - margin_y, "Document unavailable.")
            pdf.showPage()

        pdf.save()
        return converted_path

    def _wrap_text(self, text: str, max_width: float) -> List[str]:
        # Approximate characters per line based on Helvetica 11pt width (~6.2 points per char)
        if not text:
            return []
        approx_char_limit = max(int(max_width / 6.2), 40)
        lines: List[str] = []
        for paragraph in text.splitlines():
            stripped = paragraph.strip()
            if not stripped:
                lines.append("")
                continue
            wrapped = textwrap.wrap(stripped, width=approx_char_limit)
            lines.extend(wrapped if wrapped else [""])
        return lines

    def _infer_policy_metadata(
        self,
        text: str,
        *,
        fallback_title: Optional[str] = None,
        heading: Optional[str] = None,
    ) -> dict:
        """
        Attempt to enrich metadata with policy-specific hints:
        - policy_name
        - section_label
        - effective_date
        """
        metadata: dict = {}
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        policy_name = None
        section_label = None
        effective_date = None

        if heading and heading.strip():
            section_label = heading.strip()

        # Prefer an explicit heading as the policy title when available
        if heading and heading.strip():
            policy_name = heading.strip()
        else:
            if lines:
                first_line = lines[0]
                # Avoid treating obvious date/metadata lines as titles
                looks_like_date_line = bool(
                    re.search(r"^(effective(?:\s+date)?|last\s+updated|revision\s+date)[:\s]",
                             first_line, re.IGNORECASE)
                )
                if len(first_line.split()) >= 4 and not looks_like_date_line:
                    policy_name = first_line
        if not policy_name and fallback_title:
            policy_name = fallback_title

        section_match = re.search(r"(section|sec\.|article)\s+([A-Za-z0-9\.\-]+)", text, re.IGNORECASE)
        if section_match and not section_label:
            section_label = f"{section_match.group(1).title()} {section_match.group(2)}"

        date_match = re.search(
            r"(effective(?:\s+date)?|last\s+updated|revision\s+date)\s*[:\-]?\s*(?P<date>[A-Za-z]+\s+\d{1,2},\s+\d{4}|\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
            text,
            re.IGNORECASE,
        )
        if date_match:
            raw_date = date_match.group("date")
            effective_date = self._normalize_date(raw_date)

        if policy_name:
            metadata["policy_name"] = policy_name.strip()
        if section_label:
            metadata["section_label"] = section_label.strip()
        if effective_date:
            metadata["effective_date"] = effective_date
        return metadata

    def _is_arabic_corrupted(self, text: str) -> bool:
        """Detect obvious Arabic text corruption (reversed, broken characters)."""
        if not text:
            return False
        # Check for Arabic characters
        arabic_chars = re.findall(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]", text)
        if not arabic_chars or len(arabic_chars) < 10:
            return False  # Not enough Arabic to judge
        # Look for patterns that suggest corruption:
        # - Excessive isolated characters (each char appears only once)
        # - Unusual character sequences
        isolated = sum(1 for c in arabic_chars if text.count(c) == 1)
        if isolated > len(arabic_chars) * 0.3:  # More than 30% isolated
            return True
        return False

    @staticmethod
    def _normalize_date(raw_value: str) -> Optional[str]:
        raw = raw_value.strip()
        date_formats = [
            "%B %d, %Y",
            "%b %d, %Y",
            "%m/%d/%Y",
            "%m/%d/%y",
            "%d/%m/%Y",
            "%d/%m/%y",
            "%Y-%m-%d",
        ]
        for pattern in date_formats:
            try:
                parsed = datetime.strptime(raw, pattern)
                return parsed.date().isoformat()
            except ValueError:
                continue
        return raw or None
